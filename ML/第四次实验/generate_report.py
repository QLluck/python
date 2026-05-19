#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成SVM实验Word报告
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def generate_word_report(results=None):
    """生成Word实验报告"""
    
    # 创建文档
    doc = Document()
    
    # 设置标题
    title = doc.add_heading('SVM支持向量机实验报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加基本信息
    doc.add_paragraph('实验日期：2026年5月')
    doc.add_paragraph('实验课程：机器学习')
    doc.add_paragraph()
    
    # ============================================================
    # 第一部分：算法描述、数据集介绍、实验方案
    # ============================================================
    doc.add_heading('第一部分：算法描述、数据集介绍与实验方案', 1)
    
    # 1.1 算法描述
    doc.add_heading('1.1 支持向量机（SVM）算法原理', 2)
    
    p = doc.add_paragraph()
    p.add_run('支持向量机（Support Vector Machine, SVM）').bold = True
    p.add_run('是一种监督学习算法，主要用于分类和回归任务。SVM的核心思想是找到一个最优超平面，'
              '使得不同类别的样本之间的间隔最大化。')
    
    doc.add_paragraph('SVM的主要特点：', style='List Bullet')
    doc.add_paragraph('最大间隔分类器：寻找使类别间隔最大的决策边界', style='List Bullet 2')
    doc.add_paragraph('支持向量：只有靠近决策边界的样本点（支持向量）影响模型', style='List Bullet 2')
    doc.add_paragraph('核技巧：通过核函数将数据映射到高维空间，处理非线性问题', style='List Bullet 2')
    doc.add_paragraph('泛化能力强：通过最大化间隔，减少过拟合风险', style='List Bullet 2')
    
    doc.add_paragraph()
    doc.add_paragraph('常用核函数：')
    doc.add_paragraph('线性核（Linear）：K(x, y) = x^T · y，适用于线性可分数据', style='List Bullet')
    doc.add_paragraph('多项式核（Polynomial）：K(x, y) = (γx^T·y + r)^d，适用于多项式可分数据', style='List Bullet')
    doc.add_paragraph('高斯核/RBF核：K(x, y) = exp(-γ||x-y||^2)，适用于非线性数据', style='List Bullet')
    
    # 1.2 数据集介绍
    doc.add_heading('1.2 数据集介绍', 2)
    
    doc.add_heading('数据集1：Iris鸢尾花数据集', 3)
    doc.add_paragraph('来源：经典的机器学习数据集，由Fisher在1936年收集', style='List Bullet')
    doc.add_paragraph('样本数量：150个样本', style='List Bullet')
    doc.add_paragraph('特征数量：4个特征（花萼长度、花萼宽度、花瓣长度、花瓣宽度）', style='List Bullet')
    doc.add_paragraph('标签：3个类别（Setosa、Versicolour、Virginica）', style='List Bullet')
    doc.add_paragraph('特点：部分类别线性可分，适合测试线性核SVM', style='List Bullet')
    
    doc.add_heading('数据集2：乳腺癌数据集', 3)
    doc.add_paragraph('来源：威斯康星大学医院的乳腺癌诊断数据', style='List Bullet')
    doc.add_paragraph('样本数量：569个样本', style='List Bullet')
    doc.add_paragraph('特征数量：30个特征（从细胞核图像中提取）', style='List Bullet')
    doc.add_paragraph('标签：2个类别（恶性malignant、良性benign）', style='List Bullet')
    doc.add_paragraph('特点：特征维度高，非线性可分，适合测试RBF核SVM', style='List Bullet')
    
    # 1.3 实验方案
    doc.add_heading('1.3 实验方案设计', 2)
    
    doc.add_paragraph('本实验采用以下步骤：')
    doc.add_paragraph('步骤1：数据加载与探索性分析', style='List Number')
    doc.add_paragraph('步骤2：数据预处理（标准化、划分训练测试集）', style='List Number')
    doc.add_paragraph('步骤3：训练不同核函数的SVM模型（线性核、多项式核、RBF核）', style='List Number')
    doc.add_paragraph('步骤4：模型性能评估（准确率、混淆矩阵、分类报告）', style='List Number')
    doc.add_paragraph('步骤5：超参数调优（网格搜索优化C和gamma参数）', style='List Number')
    doc.add_paragraph('步骤6：结果可视化（决策边界、混淆矩阵、性能对比）', style='List Number')
    doc.add_paragraph('步骤7：实验总结与分析', style='List Number')
    
    doc.add_page_break()
    
    # ============================================================
    # 第二部分：程序清单
    # ============================================================
    doc.add_heading('第二部分：程序清单', 1)
    
    doc.add_paragraph('本实验使用Python语言实现，主要使用以下库：')
    doc.add_paragraph('numpy：数值计算', style='List Bullet')
    doc.add_paragraph('pandas：数据处理', style='List Bullet')
    doc.add_paragraph('matplotlib、seaborn：数据可视化', style='List Bullet')
    doc.add_paragraph('scikit-learn：机器学习算法库', style='List Bullet')
    doc.add_paragraph('python-docx：生成Word报告', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('完整代码请参见以下文件：')
    doc.add_paragraph('main.py：主程序入口', style='List Bullet')
    doc.add_paragraph('svm_experiment_part1.py：数据加载与预处理', style='List Bullet')
    doc.add_paragraph('svm_experiment_part2.py：模型训练与评估', style='List Bullet')
    doc.add_paragraph('svm_experiment_part3.py：超参数调优', style='List Bullet')
    doc.add_paragraph('generate_report.py：报告生成', style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('核心代码示例：').bold = True
    
    # 添加代码示例
    code_text = '''
# 导入必要的库
from sklearn.svm import SVC
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import accuracy_score, classification_report

# 数据预处理
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# 训练SVM模型（RBF核）
svm_model = SVC(kernel='rbf', random_state=42)
svm_model.fit(X_train_scaled, y_train)

# 预测和评估
y_pred = svm_model.predict(X_test_scaled)
accuracy = accuracy_score(y_test, y_pred)

# 超参数调优
param_grid = {'C': [0.1, 1, 10, 100], 'gamma': [0.001, 0.01, 0.1, 1]}
grid_search = GridSearchCV(SVC(kernel='rbf'), param_grid, cv=5)
grid_search.fit(X_train_scaled, y_train)
best_model = grid_search.best_estimator_
'''
    
    p = doc.add_paragraph(code_text)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ============================================================
    # 第三部分：实验结果
    # ============================================================
    doc.add_heading('第三部分：实验结果', 1)
    
    doc.add_heading('3.1 数据分布可视化', 2)
    
    # 添加图片
    if os.path.exists('figures/01_iris_distribution.png'):
        doc.add_paragraph('图1：Iris数据集分布图')
        doc.add_picture('figures/01_iris_distribution.png', width=Inches(6))
        doc.add_paragraph()
    
    if os.path.exists('figures/02_cancer_distribution.png'):
        doc.add_paragraph('图2：乳腺癌数据集分布图')
        doc.add_picture('figures/02_cancer_distribution.png', width=Inches(6))
        doc.add_paragraph()
    
    doc.add_heading('3.2 不同核函数的混淆矩阵', 2)
    
    if os.path.exists('figures/03_iris_confusion_matrices.png'):
        doc.add_paragraph('图3：Iris数据集不同核函数的混淆矩阵')
        doc.add_picture('figures/03_iris_confusion_matrices.png', width=Inches(6))
        doc.add_paragraph()
    
    if os.path.exists('figures/03_cancer_confusion_matrices.png'):
        doc.add_paragraph('图4：乳腺癌数据集不同核函数的混淆矩阵')
        doc.add_picture('figures/03_cancer_confusion_matrices.png', width=Inches(6))
        doc.add_paragraph()
    
    doc.add_heading('3.3 模型性能对比', 2)
    
    if os.path.exists('figures/04_model_comparison.png'):
        doc.add_paragraph('图5：不同核函数性能对比')
        doc.add_picture('figures/04_model_comparison.png', width=Inches(6))
        doc.add_paragraph()
    
    doc.add_heading('3.4 决策边界可视化', 2)
    
    if os.path.exists('figures/08_decision_boundaries.png'):
        doc.add_paragraph('图6：不同核函数的决策边界（Iris数据集）')
        doc.add_picture('figures/08_decision_boundaries.png', width=Inches(6))
        doc.add_paragraph()
    
    doc.add_heading('3.5 超参数调优结果', 2)
    
    if os.path.exists('figures/05_iris_hyperparameter_tuning.png'):
        doc.add_paragraph('图7：Iris数据集超参数调优热力图')
        doc.add_picture('figures/05_iris_hyperparameter_tuning.png', width=Inches(5))
        doc.add_paragraph()
    
    if os.path.exists('figures/05_cancer_hyperparameter_tuning.png'):
        doc.add_paragraph('图8：乳腺癌数据集超参数调优热力图')
        doc.add_picture('figures/05_cancer_hyperparameter_tuning.png', width=Inches(5))
        doc.add_paragraph()
    
    if os.path.exists('figures/07_tuning_improvement.png'):
        doc.add_paragraph('图9：调优前后性能提升对比')
        doc.add_picture('figures/07_tuning_improvement.png', width=Inches(6))
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============================================================
    # 第四部分：疑难小结
    # ============================================================
    doc.add_heading('第四部分：疑难小结与心得体会', 1)
    
    doc.add_heading('4.1 实验中遇到的问题及解决方法', 2)
    
    doc.add_heading('问题1：中文字体显示问题', 3)
    doc.add_paragraph('问题描述：在matplotlib绘图时，中文标签显示为方框', style='List Bullet')
    doc.add_paragraph('解决方法：设置plt.rcParams["font.sans-serif"]为支持中文的字体', style='List Bullet')
    doc.add_paragraph('代码：plt.rcParams["font.sans-serif"] = ["Arial Unicode MS", "SimHei"]', style='List Bullet')
    
    doc.add_heading('问题2：数据标准化的重要性', 3)
    doc.add_paragraph('问题描述：未标准化的数据导致SVM性能下降', style='List Bullet')
    doc.add_paragraph('原因分析：SVM对特征尺度敏感，不同量纲的特征会影响距离计算', style='List Bullet')
    doc.add_paragraph('解决方法：使用StandardScaler进行特征标准化', style='List Bullet')
    
    doc.add_heading('问题3：超参数选择困难', 3)
    doc.add_paragraph('问题描述：C和gamma参数的选择对模型性能影响很大', style='List Bullet')
    doc.add_paragraph('解决方法：使用GridSearchCV进行网格搜索，自动寻找最优参数组合', style='List Bullet')
    doc.add_paragraph('经验：C控制正则化强度，gamma控制核函数的影响范围', style='List Bullet')
    
    doc.add_heading('问题4：高维数据可视化', 3)
    doc.add_paragraph('问题描述：乳腺癌数据集有30个特征，无法直接可视化', style='List Bullet')
    doc.add_paragraph('解决方法：使用PCA降维到2维进行可视化', style='List Bullet')
    doc.add_paragraph('注意：降维后的可视化仅用于展示，实际模型使用全部特征', style='List Bullet')
    
    doc.add_heading('4.2 实验心得体会', 2)
    
    doc.add_paragraph('通过本次实验，我深入理解了SVM算法的原理和应用，主要收获如下：')
    doc.add_paragraph()
    
    doc.add_paragraph('1. 核函数的选择至关重要', style='List Number')
    p = doc.add_paragraph('   ', style='List Number')
    p.add_run('线性核适合线性可分数据，计算速度快；RBF核适合非线性数据，泛化能力强。'
              '实验中发现，Iris数据集使用线性核和RBF核效果相近，而乳腺癌数据集使用RBF核效果明显更好。')
    
    doc.add_paragraph('2. 数据预处理不可忽视', style='List Number')
    p = doc.add_paragraph('   ', style='List Number')
    p.add_run('特征标准化对SVM性能影响显著。标准化后，模型收敛更快，性能更稳定。')
    
    doc.add_paragraph('3. 超参数调优能显著提升性能', style='List Number')
    p = doc.add_paragraph('   ', style='List Number')
    p.add_run('通过网格搜索找到最优的C和gamma参数，模型准确率有明显提升。'
              '但需要注意计算成本，参数网格不宜过密。')
    
    doc.add_paragraph('4. 支持向量的意义', style='List Number')
    p = doc.add_paragraph('   ', style='List Number')
    p.add_run('通过可视化决策边界和支持向量，直观理解了SVM只依赖边界附近的样本点。'
              '这使得SVM对噪声数据有一定的鲁棒性。')
    
    doc.add_paragraph('5. 交叉验证的重要性', style='List Number')
    p = doc.add_paragraph('   ', style='List Number')
    p.add_run('使用交叉验证评估模型，避免了单次划分的偶然性，使评估结果更可靠。')
    
    doc.add_heading('4.3 未来改进方向', 2)
    
    doc.add_paragraph('尝试更多核函数：如Sigmoid核', style='List Bullet')
    doc.add_paragraph('处理不平衡数据：使用SMOTE或调整类权重', style='List Bullet')
    doc.add_paragraph('特征工程：尝试特征选择和特征组合', style='List Bullet')
    doc.add_paragraph('模型融合：将SVM与其他算法（如随机森林）进行集成', style='List Bullet')
    doc.add_paragraph('大规模数据：研究线性SVM和SGD优化方法', style='List Bullet')
    
    doc.add_heading('4.4 实验结论', 2)
    
    doc.add_paragraph('本实验成功实现了SVM算法在两个不同数据集上的应用，验证了以下结论：')
    doc.add_paragraph()
    doc.add_paragraph('SVM是一种强大的分类算法，特别适合中小规模、高维数据', style='List Number')
    doc.add_paragraph('核函数的选择应根据数据特性决定，RBF核适用范围最广', style='List Number')
    doc.add_paragraph('超参数调优能显著提升模型性能，但需要平衡计算成本', style='List Number')
    doc.add_paragraph('数据预处理（特别是标准化）对SVM至关重要', style='List Number')
    doc.add_paragraph('可视化有助于理解模型行为和决策过程', style='List Number')
    
    doc.add_paragraph()
    doc.add_paragraph('通过本次实验，不仅掌握了SVM的理论知识，更重要的是培养了'
                     '数据分析、模型调优和结果解释的实践能力。')
    
    # 保存文档
    doc.save('SVM实验报告.docx')
    print('Word报告已生成：SVM实验报告.docx')


if __name__ == '__main__':
    generate_word_report()
